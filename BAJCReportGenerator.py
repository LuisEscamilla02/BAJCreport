# pip install streamlit gspread oauth2client pandas reportlab python-docx matplotlib seaborn mammoth

import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import io
import gspread
import re
import textwrap
import seaborn as sns
import streamlit as st
import tempfile
import base64
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# For Staff Reports
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build

# For Campus Rep Reports
from oauth2client.service_account import ServiceAccountCredentials

#For saving spreadsheet ID 
import json

CONFIG_FILE = "config.json"

def load_config():
    try:
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}

def save_config(config: dict):
    with open(CONFIG_FILE, "w") as f:
        json.dump(config, f)
config = load_config()


# IMPORTANT CHANGE: Modify Staff Reports Likert scale options to match Campus Rep order
STAFF_LIKERT_OPTIONS = ["Strongly Agree", "Agree", "Neutral", "Disagree", "Strongly Disagree"]
STAFF_COLORS = {
    "Strongly Agree": "#1a9641",    # Green
    "Agree": "#a6d96a",             # Light green
    "Neutral": "#ffff42",           # Yellow
    "Disagree": "#fdae61",          # Intermediate orange
    "Strongly Disagree": "#d7191c"  # Red
}

# Campus Rep Reports - Fixed ordered Likert scale options (unchanged)
CAMPUS_REP_LIKERT_ORDER = ["Strongly Agree", "Agree", "Neutral", "Disagree", "Strongly Disagree"]

# Helper function for document styling
def set_cell_background(cell, hex_color):
    """Set the background color of a table cell in a Word document."""
    cell_props = cell._element.tcPr
    if cell_props is None:
        cell_props = OxmlElement('w:tcPr')
        cell._element.append(cell_props)

    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), hex_color.replace('#', ''))
    cell_props.append(cell_shading)

# Enhanced helper function for improved table styling
def style_table(table):
    """Apply consistent styling to a table"""
    # Set overall table properties
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style the header row
    header_cells = table.rows[0].cells
    for cell in header_cells:
        set_cell_background(cell, "#DDDDDD")  # Light gray background
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Ensure all data cells are center-aligned
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Helper function for document preview in Streamlit
def create_download_link(file_path, file_name):
    """Create a download link for the generated report"""
    with open(file_path, 'rb') as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{file_name}">Download {file_name}</a>'
    return href

def preview_docx(docx_path):
    """Convert DOCX to HTML for preview in Streamlit"""
    try:
        import mammoth

        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value

            # Create a styled HTML container
            styled_html = f"""
            <div style="padding: 20px; border: 1px solid #ddd; border-radius: 5px; background-color: white; color: black;">
                {html}
            </div>
            """

            return styled_html
    except ImportError:
        st.warning("DOCX preview requires additional package: pip install mammoth")
        return None
    except Exception as e:
        st.error(f"Error generating preview: {str(e)}")
        return None

# STAFF REPORT FUNCTIONS
def connect_to_google_sheets(credentials_file, spreadsheet_id):
    """Connect to Google Sheets API with service account credentials"""
    SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
    creds = Credentials.from_service_account_file(credentials_file, scopes=SCOPES)
    service = build('sheets', 'v4', credentials=creds)
    sheets = service.spreadsheets()
    return sheets

def get_sheet_names(sheets, spreadsheet_id):
    """Get all sheet names in the spreadsheet"""
    spreadsheet = sheets.get(spreadsheetId=spreadsheet_id).execute()
    sheet_names = [sheet['properties']['title'] for sheet in spreadsheet['sheets']]
    return sheet_names

def get_sheet_data(sheets, spreadsheet_id, sheet_name):
    """Retrieve data from specified sheet"""
    range_name = f"'{sheet_name}'!A1:AAA1000"  # Adjust range as needed
    result = sheets.values().get(
        spreadsheetId=spreadsheet_id,
        range=range_name
    ).execute()
    values = result.get('values', [])
    return values

def clean_column_names(df):
    """Clean column names by extracting staff name and skill category"""
    cleaned_columns = []
    for col in df.columns:
        if '[' in col and ']' in col:
            # Extract the staff name and skill category
            match = re.match(r"(.*?)\s*\[(.*?)\]", col)
            if match:
                staff_name = match.group(1).strip()
                skill = match.group(2).strip()
                cleaned_columns.append(f"{staff_name} - {skill}")
            else:
                cleaned_columns.append(col)
        else:
            cleaned_columns.append(col)
    df.columns = cleaned_columns
    return df

def calculate_likert_counts(df, staff_name):
    """
    Calculate the raw counts of Likert responses for each skill metric
    for the given staff member.
    """


    # Get columns associated with this staff member (expecting headers like "Name - Skill")
    staff_columns = [col for col in df.columns if staff_name in col]
    counts_data = {}

    # For each column (skill metric), count responses
    for col in staff_columns:
        # Extract skill metric name from the column header
        skill = col.split(" - ")[-1]
        counts = {option: 0 for option in STAFF_LIKERT_OPTIONS}
        
        for val in df[col]:
            # Account for possible inconsistency in response format
            for option in STAFF_LIKERT_OPTIONS:
                if str(val).strip() == option:
                    counts[option] += 1
                    break

        counts_data[skill] = counts
    return counts_data

def create_clustered_bar_chart_staff(counts_data, staff_name):
    """
    Create a clustered bar chart with raw counts for staff report and return as bytes for embedding in Word
    """
    skills = list(counts_data.keys())
    n_skills = len(skills)
    n_options = len(STAFF_LIKERT_OPTIONS)

    # Create the figure and axis
    fig, ax = plt.subplots(figsize=(7, 4))

    # Positions for the clusters on the x-axis
    indices = np.arange(n_skills)
    bar_width = 0.15  # Slightly wider to accommodate fewer bars

    # Create stacked bar chart using raw counts
    for i, option in enumerate(STAFF_LIKERT_OPTIONS):
        # Get raw count values for this option across all skills
        counts = [counts_data[skill][option] for skill in skills]
        # Compute bar positions for this option within each cluster
        positions = indices + i * bar_width - (n_options/2)*bar_width + bar_width/2
        ax.bar(positions, counts, width=bar_width, color=STAFF_COLORS[option], label=option)

    ax.set_xlabel("Skill Metrics")
    ax.set_ylabel("Count")
    ax.set_title(f"{staff_name} - Likert Scale Response Distribution")
    ax.set_xticks(indices)
    wrapped_labels = [textwrap.fill(skill, width=13) for skill in skills]  # Adjust width as needed
    ax.set_xticklabels(wrapped_labels, rotation=0, ha='center', fontsize=8)
    
    # Position legend for better readability
    ax.legend(loc='upper right', bbox_to_anchor=(1, 1), ncol=1, fontsize=8)
    ax.grid(axis='y', linestyle='--', alpha=0.7)

    plt.tight_layout()

    # Save the chart as bytes to embed in Word
    img_bytes = io.BytesIO()
    plt.savefig(img_bytes, format='png', dpi=300)
    img_bytes.seek(0)
    plt.close(fig)

    return img_bytes

def generate_staff_report_docx(staff_name, df, output_folder="reports"):
    """
    Generate a Word (.docx) report for the given staff member
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    output_file = os.path.join(output_folder, f"{staff_name.replace(' ', '_')}_Report.docx")

    # Calculate the raw counts of responses for the staff member
    counts_data = calculate_likert_counts(df, staff_name)

    # Create a new Word document
    doc = Document()

    # Add title
    title = doc.add_heading(f"Staff Report: {staff_name}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add chart
    chart_bytes = create_clustered_bar_chart_staff(counts_data, staff_name)
    doc.add_picture(chart_bytes, width=Inches(6.5))

    # Add caption for chart
    chart_caption = doc.add_paragraph("Figure 1: Likert Scale Response Count Distribution")
    chart_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_caption.style = 'Caption'

    # Add table header
    doc.add_heading("Response Summary", level=2)

    # Create summary table
    table = doc.add_table(rows=len(counts_data)+1, cols=len(STAFF_LIKERT_OPTIONS)+1)
    style_table(table)  # Apply consistent styling

    # Set header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Skill Category"
    for i, option in enumerate(STAFF_LIKERT_OPTIONS):
        header_cells[i+1].text = option

    # Fill in data rows with raw counts
    for i, (skill, counts) in enumerate(counts_data.items()):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = skill

        for j, option in enumerate(STAFF_LIKERT_OPTIONS):
            row_cells[j+1].text = str(counts[option])  # Raw count as string

    # Add space after table
    doc.add_paragraph()

    # Save the document
    doc.save(output_file)

    return output_file

def process_staff_evaluation_docx(credentials_file, spreadsheet_id, sheet_name, staff_name):
    """
    Process the staff evaluation data and generate a DOCX report
    """
    # Connect to Google Sheets and get data
    sheets_api = connect_to_google_sheets(credentials_file, spreadsheet_id)
    data = get_sheet_data(sheets_api, spreadsheet_id, sheet_name)

    if not data:
        st.error(f"No data found in sheet '{sheet_name}'")
        return {}

    df = pd.DataFrame(data[1:], columns=data[0])
    df = clean_column_names(df)

    docx_file = generate_staff_report_docx(staff_name, df)
    st.success(f"Generated report for {staff_name}: {docx_file}")
    return {staff_name: docx_file}

# CAMPUS REP REPORT FUNCTIONS
def list_sheet_names(creds_file, sheet_id):
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds = ServiceAccountCredentials.from_json_keyfile_name(creds_file, scope)
    client = gspread.authorize(creds)
    workbook = client.open_by_key(sheet_id)
    worksheets = workbook.worksheets()
    sheet_names = [sheet.title for sheet in worksheets]
    return sheet_names

def get_data(creds_file, sheet_id, sheet_name):
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds = ServiceAccountCredentials.from_json_keyfile_name(creds_file, scope)
    client = gspread.authorize(creds)
    sheet = client.open_by_key(sheet_id).worksheet(sheet_name)

    # Instead of using get_all_records, we'll get all values and handle them directly
    all_values = sheet.get_all_values()
    if not all_values:
        return pd.DataFrame()

    # Get headers (first row)
    headers = all_values[0]

    # Check for duplicate headers and make them unique if needed
    unique_headers = []
    header_count = {}

    for header in headers:
        if header in header_count:
            header_count[header] += 1
            unique_headers.append(f"{header}_{header_count[header]}")
        else:
            header_count[header] = 0
            unique_headers.append(header)

    # Create DataFrame with the unique headers
    df = pd.DataFrame(all_values[1:], columns=unique_headers)

    return df

def generate_clustered_bar_graph_campus_rep(df_staff, likert_columns, question_labels=None):
    """
    Generates a clustered bar graph showing raw counts for each Likert question.
    Styled to match the staff member graph style.
    Returns image bytes for embedding in Word.
    """
    num_questions = len(likert_columns)
    if question_labels is None:
        question_labels = [f"Q{i+1}" for i in range(num_questions)]

    # Create color mapping matching the staff report colors
    campus_rep_colors = {
        "Strongly Agree": "#1a9641",    # Green
        "Agree": "#a6d96a",             # Light green
        "Neutral": "#ffff42",           # Yellow
        "Disagree": "#fdae61",          # Intermediate orange
        "Strongly Disagree": "#d7191c"  # Red
    }

    # Prepare data structure: a dict mapping each Likert option to a list of raw counts per question
    summary_data = {option: [] for option in CAMPUS_REP_LIKERT_ORDER}

    for i, col in enumerate(likert_columns):
        responses = df_staff[col].dropna()
        value_counts = responses.value_counts().to_dict()
        
        for option in CAMPUS_REP_LIKERT_ORDER:
            count = value_counts.get(option, 0)
            summary_data[option].append(count)

    # Create the figure and axis with matching dimensions to staff report
    fig, ax = plt.subplots(figsize=(7, 4))  # Match staff report dimensions

    # Positions for the clusters on the x-axis
    indices = np.arange(num_questions)
    bar_width = 0.15  # Match staff report width

    # Plot bars with consistent style
    for i, option in enumerate(CAMPUS_REP_LIKERT_ORDER):
        # Get raw count values for this option across all questions
        counts = summary_data[option]
        # Compute bar positions for this option within each cluster
        positions = indices + i * bar_width - (len(CAMPUS_REP_LIKERT_ORDER)/2)*bar_width + bar_width/2
        ax.bar(positions, counts, width=bar_width, color=campus_rep_colors[option], label=option)

    # Set labels and title with consistent styling
    staff_member = df_staff["Name of Campus Representative"].iloc[0] if not df_staff.empty else "Unknown"
    ax.set_xlabel("Question Metrics")
    ax.set_ylabel("Count")
    ax.set_title(f"{staff_member} - Likert Scale Response Distribution")

    ax.set_xticks(indices)
    # Use textwrap to wrap long labels as in staff report
    wrapped_labels = [textwrap.fill(label, width=13) for label in question_labels]
    ax.set_xticklabels(wrapped_labels, rotation=0, ha='center', fontsize=8)

    ax.legend(loc='upper right', bbox_to_anchor=(1, 1), ncol=1, fontsize=8)  # Match staff legend
    ax.grid(axis='y', linestyle='--', alpha=0.7)  # Match staff grid style

    plt.tight_layout()

    # Save the chart as bytes to embed in Word
    img_bytes = io.BytesIO()
    plt.savefig(img_bytes, format='png', dpi=300)
    img_bytes.seek(0)
    plt.close(fig)

    return img_bytes

def generate_docx_report(df, staff_member, output_folder="reports"):
    """
    Generates a Word (.docx) report for the given campus rep,
    including a clustered bar graph, a summary table, and a comments section.
    Saves to the same 'reports' folder as staff reports for consistency.
    """
    # Create reports directory if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Create a safe file name that includes the staff member's name
    safe_name = "".join(c for c in staff_member if c.isalnum() or c in ['_']).strip()
    output_filename = os.path.join(output_folder, f"campus_rep_report_{safe_name}.docx")

    # Filter data for the selected staff member
    df_staff = df[df["Name of Campus Representative"] == staff_member]
    if df_staff.empty:
        st.error(f"No data available for staff member: {staff_member}")
        return None

    # List of Likert-scale question columns
    likert_columns = [
        "For each statement below, please select the option that best represents your opinion [My Campus Rep is accessible, respectful, and responsive to my needs for support.]",
        "For each statement below, please select the option that best represents your opinion [I have received regular, clear communication from my Campus Rep.]",
        "For each statement below, please select the option that best represents your opinion [Reflection sessions (during Saturday trainings, as well as on-campus) with my Campus Rep have been useful.]",
        "For each statement below, please select the option that best represents your opinion [I would have liked more time to reflect upon my JusticeCorps experiences with my peers.]",
        "For each statement below, please select the option that best represents your opinion [My Campus Rep has been a good resource.]"
    ]

    # Abbreviated question names for readability in tables and charts
    question_labels = [
        "Accessible & Responsive",
        "Clear Communication",
        "Useful Reflection Sessions",
        "More Peer Reflection Time",
        "Good Resource"
    ]

    # Create a new Word document
    doc = Document()

    # Add title
    title = doc.add_heading(f"Campus Rep Report: {staff_member}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Generate and add chart - now passing question_labels for better readability
    chart_bytes = generate_clustered_bar_graph_campus_rep(df_staff, likert_columns, question_labels)
    doc.add_picture(chart_bytes, width=Inches(6.5))

    # Add caption for chart
    chart_caption = doc.add_paragraph("Figure 1: Likert Scale Response Count Distribution")
    chart_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_caption.style = 'Caption'

    # Add table header - match staff report header style
    doc.add_heading("Response Summary", level=2)  # Match staff report heading

    # Create summary table
    table = doc.add_table(rows=len(question_labels)+1, cols=len(CAMPUS_REP_LIKERT_ORDER)+1)
    style_table(table)  # Apply consistent styling

    # Set header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Question"  # Could also use "Skill Category" to match staff report exactly
    for i, option in enumerate(CAMPUS_REP_LIKERT_ORDER):
        header_cells[i+1].text = option

    # Fill in data rows with raw counts
    for i, (col, label) in enumerate(zip(likert_columns, question_labels)):
        responses = df_staff[col].dropna()
        row_cells = table.rows[i+1].cells
        row_cells[0].text = label  # Use short label without Q# prefix for cleaner look

        # Count responses for each Likert option
        value_counts = responses.value_counts().to_dict()
        for j, option in enumerate(CAMPUS_REP_LIKERT_ORDER):
            count = value_counts.get(option, 0)
            row_cells[j+1].text = str(count)  # Display raw count
            paragraph = row_cells[j+1].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Comments section
    doc.add_heading("Comments", level=2)
    comments = df_staff["Comments on Campus Representative Support"].dropna().tolist()
    if comments:
        for comment in comments:
            doc.add_paragraph(comment, style='List Bullet')
    else:
        doc.add_paragraph("No comments available.")

    # Save the document
    doc.save(output_filename)
    st.success(f"DOCX report saved as {output_filename}")
    return output_filename

# STREAMLIT APP INTERFACE
def streamlit_app():
    st.set_page_config(
        page_title="Campus Rep & Staff Reporting System",
        page_icon="📊",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    st.title("Campus Rep & Staff Reporting System")
    st.markdown("---")

    # Create reports directory if it doesn't exist
    if not os.path.exists("reports"):
        os.makedirs("reports")

    # Initialize session state for spreadsheet_id
    if 'spreadsheet_id' not in st.session_state:
        st.session_state.spreadsheet_id = ""

    with st.sidebar:
        st.header("Configuration")
        report_type = st.radio(
            "Select Report Type",
            ["Campus Rep Report", "Staff Report"]
        )
        st.markdown("---")

        # File uploader for credentials
        uploaded_file = st.file_uploader(
            "Upload Google API credentials JSON", type=['json']
        )
        if uploaded_file is not None:
            # Save the uploaded file to a temporary location
            with tempfile.NamedTemporaryFile(delete=False, suffix='.json') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                credentials_path = tmp_file.name
            st.session_state.credentials_file = credentials_path

        # Google Sheet ID input bound to session_state
        st.text_input(
            "Google Sheet ID",
            key="spreadsheet_id"
        )

        # Button to save Spreadsheet ID in session
        if st.button("💾 Save Spreadsheet ID"):
            if st.session_state.spreadsheet_id.strip():
                st.success("Spreadsheet ID saved in session! ✅")
            else:
                st.warning("Cannot save empty Spreadsheet ID.")

    # Only proceed when credentials and sheet ID are provided
    if st.session_state.get('credentials_file') and st.session_state.spreadsheet_id:
        credentials_path = st.session_state.credentials_file
        sheet_id = st.session_state.spreadsheet_id

        if report_type == "Staff Report":
            st.header("Staff Report Generation")
            try:
                # Connect to Google Sheets to get available sheet names
                sheets_api = connect_to_google_sheets(credentials_path, sheet_id)
                sheet_names = get_sheet_names(sheets_api, sheet_id)

                if not sheet_names:
                    st.error("No sheets found in the spreadsheet.")
                else:
                    selected_sheet = st.selectbox("Select a sheet", sheet_names)
                    staff_name = st.text_input("Enter staff name to process")

                    if staff_name and st.button("Generate Staff Report"):
                        with st.spinner(f"Generating report for {staff_name}..."):
                            result = process_staff_evaluation_docx(
                                credentials_path,
                                sheet_id,
                                selected_sheet,
                                staff_name
                            )
                            if staff_name in result:
                                st.session_state.report_generated = True
                                st.session_state.report_path = result[staff_name]
                                st.session_state.report_type = "staff"
                                st.success(f"Report for {staff_name} generated successfully!")
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

        else:  # Campus Rep Report
            st.header("Campus Rep Report Generation")
            try:
                sheet_names = list_sheet_names(credentials_path, sheet_id)
                selected_sheet = st.selectbox("Select a sheet", sheet_names)

                with st.spinner("Fetching data from Google Sheet..."):
                    df = get_data(credentials_path, sheet_id, selected_sheet)

                unique_staff = df["Name of Campus Representative"].dropna().unique().tolist()
                staff_member = st.selectbox("Select Campus Representative", unique_staff)

                if st.button("Generate Campus Rep Report"):
                    with st.spinner(f"Generating report for {staff_member}..."):
                        output_file = generate_docx_report(df, staff_member)
                        if output_file:
                            st.session_state.report_generated = True
                            st.session_state.report_path = output_file
                            st.session_state.report_type = "campus_rep"
                            st.success(f"Report for {staff_member} generated successfully!")
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

    # Show report preview if generated
    if st.session_state.get('report_generated') and st.session_state.get('report_path'):
        st.markdown("---")
        st.header("Report Preview")

        col1, col2 = st.columns([3, 1])
        with col2:
            st.markdown("### Actions")
            st.markdown(
                create_download_link(
                    st.session_state.report_path,
                    os.path.basename(st.session_state.report_path)
                ),
                unsafe_allow_html=True
            )
            if st.button("Generate Another Report"):
                st.session_state.report_generated = False
                st.session_state.report_path = None
                st.rerun()

        with col1:
            st.markdown("### Document Preview")
            try:
                html_preview = preview_docx(st.session_state.report_path)
                if html_preview:
                    st.markdown(html_preview, unsafe_allow_html=True)
                else:
                    st.info("Preview not available. Install 'mammoth' package for preview: pip install mammoth")
            except Exception as e:
                st.warning(f"Could not generate preview: {str(e)}")
                st.info("Download the report to view it in Microsoft Word or compatible application.")

if __name__ == "__main__":
    streamlit_app()
